# app/api/misc.py
# -*- coding: utf-8 -*-
"""杂项接口:任务进度 / Token 用量 / 导出。

GET /api/jobs/{job_id}                任务进度(异步生成轮询)
GET /api/usage                        Token 用量汇总
GET /api/projects/{id}/export/txt     整本导出 txt
GET /api/projects/{id}/export/epub    整本导出 epub
GET /api/projects/{id}/export/md      整本导出 Markdown
GET /api/projects/{id}/export/docx    整本导出 Word(中文排版)
GET /api/projects/{id}/export/chapters-zip  按章拆成多个 txt 打包 zip
"""
from __future__ import annotations

import io
import zipfile
from html import escape
from urllib.parse import quote

from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import Response
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.auth import assert_project_owner, current_user_id, get_current_user
from app.db.models import Chapter, LlmUsage, Outline, Project
from app.db.session import get_db
from app.jobs import get_job, list_for_user, list_running

router = APIRouter(tags=["misc"], dependencies=[Depends(get_current_user)])


@router.get("/api/jobs")
async def my_jobs(all: bool = False):
    """当前用户的后台任务(全局任务中心数据源)。all=true 时含近期已完成/失败的。"""
    items = list_for_user(current_user_id.get(), running_only=not all)
    return {
        "jobs": [
            {
                "job_id": jid,
                "kind": job["kind"],
                "status": job["status"],
                "stage": job["stage"],
                "error": job.get("error"),
            }
            for jid, job in items
        ]
    }


@router.get("/api/jobs/{job_id}")
async def job_status(job_id: str):
    job = get_job(job_id)
    # 归属校验:非本人的任务按"不存在"处理,不泄露 job 存在性
    if job is None or job.get("owner_id") != current_user_id.get():
        raise HTTPException(status_code=404, detail="任务不存在或已被清理")
    job.pop("owner_id", None)  # 内部字段,不下发
    return job


@router.get("/api/projects/{project_id}/running-jobs")
async def project_running_jobs(project_id: int, db: Session = Depends(get_db)):
    """本项目正在运行的后台任务。切走页面再回来时,前端据此重新接上轮询。"""
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status_code=404, detail="项目不存在")
    assert_project_owner(project)
    jobs = []
    # 章节类按前缀匹配(kind 带章号);架构/蓝图必须精确匹配,防 project 1 命中 11
    for prefix in (f"chapter-{project_id}-", f"re-extract-{project_id}-"):
        for jid, job in list_running(prefix):
            jobs.append({"job_id": jid, "kind": job["kind"], "stage": job["stage"]})
    for exact in (f"architecture-{project_id}", f"blueprint-{project_id}"):
        for jid, job in list_running(exact):
            if job["kind"] == exact:
                jobs.append({"job_id": jid, "kind": job["kind"], "stage": job["stage"]})
    return {"jobs": jobs}


@router.get("/api/usage")
async def usage_summary(db: Session = Depends(get_db)):
    """Token 用量汇总(总量 + 按模型)——只统计当前用户。"""
    uid = current_user_id.get()
    rows = (
        db.query(
            LlmUsage.model,
            func.count(LlmUsage.id),
            func.sum(LlmUsage.prompt_tokens),
            func.sum(LlmUsage.completion_tokens),
        )
        .filter(LlmUsage.user_id == uid)
        .group_by(LlmUsage.model)
        .all()
    )
    by_model = [
        {
            "model": m, "calls": c,
            "prompt_tokens": int(p or 0), "completion_tokens": int(o or 0),
        }
        for m, c, p, o in rows
    ]
    return {
        "total_calls": sum(x["calls"] for x in by_model),
        "total_prompt_tokens": sum(x["prompt_tokens"] for x in by_model),
        "total_completion_tokens": sum(x["completion_tokens"] for x in by_model),
        "by_model": by_model,
    }


def _book(db: Session, project_id: int) -> tuple[Project, list[tuple[str, str]]]:
    """取书名与 [(章标题, 正文)],按章号排序,只含有正文的章。"""
    project = db.get(Project, project_id)
    if project is None:
        raise HTTPException(status_code=404, detail="项目不存在")
    assert_project_owner(project)
    titles = {
        o.chapter_number: o.title
        for o in db.query(Outline).filter(Outline.project_id == project_id)
    }
    chapters = (
        db.query(Chapter)
        .filter(Chapter.project_id == project_id, Chapter.final_content != "")
        .order_by(Chapter.chapter_number)
        .all()
    )
    if not chapters:
        raise HTTPException(status_code=400, detail="还没有已定稿的章节")
    items = [
        (f"第{c.chapter_number}章 {titles.get(c.chapter_number, '')}".strip(), c.final_content)
        for c in chapters
    ]
    return project, items


def _disposition(project: Project, ext: str, suffix: str = "") -> str:
    """生成下载头:文件名用书名(剔除非法字符),UTF-8 百分号编码(RFC 5987),
    并附一个纯 ASCII 回退名(老浏览器用)。书名为空时退回项目 id。
    suffix 可选,拼在书名后用于区分同书的不同导出(如"(分章)")。"""
    base = (project.title or "").strip()
    safe = "".join(c for c in base if c not in '\\/:*?"<>|').strip() or str(project.id)
    encoded = quote(f"{safe}{suffix}.{ext}", safe="")
    return f"attachment; filename=\"{project.id}.{ext}\"; filename*=UTF-8''{encoded}"


@router.get("/api/projects/{project_id}/export/txt")
async def export_txt(project_id: int, db: Session = Depends(get_db)):
    project, items = _book(db, project_id)
    parts = [f"《{project.title}》\n"]
    for title, text in items:
        parts.append(f"\n\n{title}\n\n{text}")
    data = "\n".join(parts).encode("utf-8")
    return Response(
        content=data,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": _disposition(project, "txt")},
    )


@router.get("/api/projects/{project_id}/export/epub")
async def export_epub(project_id: int, db: Session = Depends(get_db)):
    """最小可用 epub(纯标准库 zip 打包,无外部依赖)。"""
    project, items = _book(db, project_id)

    def xhtml(title: str, body: str) -> str:
        paras = "".join(
            f"<p>{escape(p.strip())}</p>" for p in body.splitlines() if p.strip()
        )
        return (
            '<?xml version="1.0" encoding="utf-8"?>\n'
            '<html xmlns="http://www.w3.org/1999/xhtml"><head>'
            f"<title>{escape(title)}</title></head><body>"
            f"<h2>{escape(title)}</h2>{paras}</body></html>"
        )

    manifest, spine, files = [], [], []
    for i, (title, text) in enumerate(items, 1):
        fn = f"chap{i}.xhtml"
        manifest.append(
            f'<item id="c{i}" href="{fn}" media-type="application/xhtml+xml"/>'
        )
        spine.append(f'<itemref idref="c{i}"/>')
        files.append((fn, xhtml(title, text)))

    nav_lis = "".join(
        f'<li><a href="chap{i}.xhtml">{escape(t)}</a></li>'
        for i, (t, _) in enumerate(items, 1)
    )
    nav = (
        '<?xml version="1.0" encoding="utf-8"?>'
        '<html xmlns="http://www.w3.org/1999/xhtml" xmlns:epub="http://www.idpf.org/2007/ops">'
        '<head><title>目录</title></head><body><nav epub:type="toc"><ol>'
        + nav_lis + "</ol></nav></body></html>"
    )
    opf = (
        '<?xml version="1.0" encoding="utf-8"?>'
        '<package xmlns="http://www.idpf.org/2007/opf" version="3.0" unique-identifier="uid">'
        '<metadata xmlns:dc="http://purl.org/dc/elements/1.1/">'
        f'<dc:identifier id="uid">jarvis-write-{project.id}</dc:identifier>'
        f"<dc:title>{escape(project.title)}</dc:title>"
        "<dc:language>zh-CN</dc:language>"
        '<meta property="dcterms:modified">2026-01-01T00:00:00Z</meta>'
        "</metadata><manifest>"
        '<item id="nav" href="nav.xhtml" media-type="application/xhtml+xml" properties="nav"/>'
        + "".join(manifest)
        + "</manifest><spine>" + "".join(spine) + "</spine></package>"
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr("mimetype", "application/epub+zip", zipfile.ZIP_STORED)
        z.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0"?><container version="1.0" '
            'xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="OEBPS/content.opf" '
            'media-type="application/oebps-package+xml"/></rootfiles></container>',
        )
        z.writestr("OEBPS/content.opf", opf)
        z.writestr("OEBPS/nav.xhtml", nav)
        for fn, content in files:
            z.writestr(f"OEBPS/{fn}", content)

    return Response(
        content=buf.getvalue(),
        media_type="application/epub+zip",
        headers={"Content-Disposition": _disposition(project, "epub")},
    )


@router.get("/api/projects/{project_id}/export/md")
async def export_md(project_id: int, db: Session = Depends(get_db)):
    """整本导出 Markdown:书名做一级标题,每章做二级标题。"""
    project, items = _book(db, project_id)
    parts = [f"# 《{project.title}》\n"]
    for title, text in items:
        paras = "\n\n".join(p.strip() for p in text.splitlines() if p.strip())
        parts.append(f"\n## {title}\n\n{paras}\n")
    data = "\n".join(parts).encode("utf-8")
    return Response(
        content=data,
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": _disposition(project, "md")},
    )


@router.get("/api/projects/{project_id}/export/chapters-zip")
async def export_chapters_zip(project_id: int, db: Session = Depends(get_db)):
    """按章拆成多个 txt 打包 zip,方便分章发布。文件名带章号便于排序。"""
    project, items = _book(db, project_id)
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for i, (title, text) in enumerate(items, 1):
            # 文件名:001 第1章 标题.txt(章号补零排序;剔除非法文件名字符)
            safe = "".join(c for c in title if c not in '\\/:*?"<>|').strip()
            z.writestr(f"{i:03d} {safe}.txt", f"{title}\n\n{text}")
    return Response(
        content=buf.getvalue(),
        media_type="application/zip",
        headers={"Content-Disposition": _disposition(project, "zip", "(分章)")},
    )


def _set_cjk(run, font_name: str = "宋体") -> None:
    """给 run 同时设 ascii 与 eastAsia 字体,避免中文走 Word 回退字体导致排版不一致。"""
    from docx.oxml.ns import qn

    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), font_name)


@router.get("/api/projects/{project_id}/export/docx")
async def export_docx(project_id: int, db: Session = Depends(get_db)):
    """整本导出 Word:中文排版(宋体、1.5倍行距、正文首行缩进2字符)。"""
    from docx import Document
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    project, items = _book(db, project_id)
    doc = Document()

    # 全局正文样式:宋体 + 小四 + 1.5倍行距
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal._element.get_or_add_rPr().get_or_add_rFonts().set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "宋体"
    )
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # 书名(居中大标题)
    t = doc.add_heading(level=0)
    t.alignment = 1  # 居中
    trun = t.add_run(f"《{project.title}》")
    _set_cjk(trun, "微软雅黑")

    for title, text in items:
        h = doc.add_heading(level=1)
        hrun = h.add_run(title)
        _set_cjk(hrun, "微软雅黑")
        for para in text.splitlines():
            para = para.strip()
            if not para:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)  # 首行缩进约2字符
            run = p.add_run(para)
            _set_cjk(run, "宋体")

    out = io.BytesIO()
    doc.save(out)
    return Response(
        content=out.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": _disposition(project, "docx")},
    )
